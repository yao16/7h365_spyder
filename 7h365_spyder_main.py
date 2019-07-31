# -*- coding: utf-8 -*-
import requests
import re
from bs4 import BeautifulSoup as bs

from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
from docx.oxml.ns import qn
#http://www.7h365.com/plugin.php?id=votes:view&appid=4815中有特殊字符
url = 'http://www.7h365.com/plugin.php?id=votes%3Alist&actid=180page%3D4&page='
person_page_url = 'http://www.7h365.com/plugin.php?id=votes:view&appid='

person_list_url = 'http://www.7h365.com/plugin.php?id=votes%3Alist&actid='
'''
vote_list_current_page = 1
vote_list_page_count = 4
while vote_list_current_page <= int(vote_list_page_count):
    current_page = requests.get(url + str(vote_list_current_page))
    current_page.encoding = 'utf-8'
    soup = bs(current_page.text, 'html.parser')
    person_id_urls = soup.select('.inside_right_photo')[0].select('.c_b_t_ls')
    print(person_id_urls)
    for person_id_url in person_id_urls:
        person_id = re.findall('\d+', person_id_url.attrs.get('href'))[0]
        person_page = requests.get(person_page_url + person_id)
        person_page.encoding = 'utf-8'
        person_page_soup = bs(person_page.text, 'html.parser')
        person_infos = person_page_soup.select('.clearfix')[2].select('li')
        for person_info in person_infos:
            print(person_info.text)
        if len(person_page_soup.select('.jianjie')) > 0 :
            print('自荐综述')
            for item in person_page_soup.select('.jianjie'):
                print(item.text)
        if len(person_page_soup.select('.art-title')) > 0:
            print('参评报告')
            for item in person_page_soup.select('.art-title'):
                print(item.text)
        print('**************************************************************')
    vote_list_current_page += 1
'''



def load_person_info(actid, vote_list_page_count):
    document = Document()
    vote_list_current_page = 1
    total_count = 0
    while vote_list_current_page <= int(vote_list_page_count):
        current_page = requests.get(person_list_url + str(actid) + 'page%3D4&page=' + str(vote_list_current_page))
        print(person_list_url + str(actid) + 'page%3D4&page=' + str(vote_list_current_page))
        current_page.encoding = 'utf-8'
        soup = bs(current_page.text, 'html.parser')
        person_id_urls = soup.select('.inside_right_photo')[0].select('.c_b_t_ls')
        print(person_id_urls)
        for person_id_url in person_id_urls:
            p = document.add_paragraph()

            person_id = re.findall('\d+', person_id_url.attrs.get('href'))[0]
            person_page = requests.get(person_page_url + person_id)
            person_page.encoding = 'utf-8'
            person_page_soup = bs(person_page.text, 'html.parser')
            person_infos = person_page_soup.select('.clearfix')[2].select('li')
            for person_info in person_infos:
                print(person_info.text)
                run = p.add_run(person_info.text + '\n')
                run.font.size = Pt(16)
                run.font.name = u'宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                run.bold = True

            if len(person_page_soup.select('.jianjie')) > 0:
                print('自荐综述')
                run = p.add_run(u'自荐综述\n')
                run.font.size = Pt(12)
                run.font.name = u'微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
                for item in person_page_soup.select('.jianjie'):
                    print(item.text)
                    if item.text != None and len(item.text) > 0:
                        run = p.add_run(item.text)
                        run.font.size = Pt(10.5)
                        run.font.name = u'微软雅黑'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
            if len(person_page_soup.select('.art-title')) > 0:
                print('参评报告')
                run = p.add_run(u'参评报告\n')
                run.font.size = Pt(12)
                run.font.name = u'微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
                for item in person_page_soup.select('.art-title'):
                    print(item.text)
                    if item.text != None and len(item.text) > 0:
                        run = p.add_run(re.sub('', '', item.text))
                        run.font.size = Pt(10.5)
                        run.font.name = u'微软雅黑'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
            print('**************************************************************')
            run = p.add_run(u'**************************************************************')
            run.font.size = Pt(16)
            run.font.name = u'宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            total_count += 1
        vote_list_current_page += 1
    document.save('c:/Users/yao16/Desktop/test180.docx')
    print(total_count)

#抓取“最佳宏观策略分析师”actid = 180
#load_person_info(180, 4)

#抓取“最佳金融期货分析师”actid = 181
#load_person_info(181, 4)

#抓取“最佳工业品期货分析师”actid = 182
#load_person_info(182, 19)

#抓取“最佳农副产品期货分析师”actid = 183 //有问题
load_person_info(183, 4)

#抓取“最佳金融量化策略工程师”actid = 184
#load_person_info(184, 1)

#抓取“最佳期权分析师”actid = 185
#load_person_info(185, 2)