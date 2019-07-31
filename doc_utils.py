from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt


document = Document()
document.add_paragraph('my paragraph')

document.styles['Normal'].font.name = u'黑体'
p = document.add_paragraph()

run = p.add_run(u'我添加的段落文字 ')
run.font.color.rgb = RGBColor(54, 95, 145)
run.font.size = Pt(36)

document.save('c:/Users/yaoxiao/Desktop/test1.docx')




